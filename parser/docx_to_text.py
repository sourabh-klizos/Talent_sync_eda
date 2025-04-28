import argparse
import re

from docx import Document
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph
from thefuzz import fuzz


class DocxExtractor:
    @staticmethod
    def iter_block_items(parent):
        """
        Yield paragraphs and tables in document order with improved handling of complex structures.
        Handles detection of paragraphs in tables robustly.
        """
        if not hasattr(parent, "element") or not hasattr(parent.element, "body"):
            return

        # Collect all paragraphs that are inside tables for accurate tracking
        table_paragraphs = set()
        try:
            # Use XPath to find all paragraphs within table cells
            for tbl in parent.element.body.findall(".//w:tbl", namespaces={"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}):
                for cell in tbl.findall(".//w:tc", namespaces={"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}):
                    for para in cell.findall(".//w:p", namespaces={"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}):
                        table_paragraphs.add(para)
        except Exception as e:
            print(f"Warning: Error collecting table paragraphs: {e}")

        # Track processed tables to avoid duplicates from nested tables
        processed_tables = set()

        # Process document in order
        try:
            for child in parent.element.body:
                if child.tag == qn("w:tbl"):
                    # Skip tables we've already processed
                    if child in processed_tables:
                        continue
                    processed_tables.add(child)
                    yield Table(child, parent)
                elif child.tag == qn("w:p"):
                    paragraph = Paragraph(child, parent)
                    # Mark if this paragraph is within a table
                    paragraph._in_table_context = child in table_paragraphs
                    yield paragraph
        except Exception as e:
            print(f"Warning: Error iterating document blocks: {e}")
            # Try a more basic approach if the sophisticated one fails
            if hasattr(parent, "paragraphs"):
                for para in parent.paragraphs:
                    yield para
            if hasattr(parent, "tables"):
                for table in parent.tables:
                    yield table

    @staticmethod
    def get_content_fingerprint(text):
        """
        Create a robust fingerprint that can identify similar content despite minor differences.
        """
        if not text:
            return ""

        # Normalize text
        text = text.lower().strip()
        # Remove extra whitespace
        text = re.sub(r"\s+", " ", text)
        # Remove common punctuation that doesn't affect meaning
        text = re.sub(r'[,.;:!?"\']', "", text)

        if len(text) <= 20:
            return text

        # Use multiple aspects of the text to create a robust fingerprint
        words = text.split()
        word_count = len(words)

        # Get beginning, middle and end samples
        start = text[:15]
        middle = text[len(text) // 2 - 7 : len(text) // 2 + 7] if len(text) > 30 else ""
        end = text[-15:]

        # Include first letters of first 8 words (if available)
        word_initials = "".join(word[0] for word in words[:8] if word)

        return f"{start}_{word_count}_{middle}_{end}_{word_initials}"

    @staticmethod
    def extract_cell_text(cell):
        """
        Extract text from a cell with proper handling of paragraph structure.
        """
        parts = []
        try:
            for paragraph in cell.paragraphs:
                if not paragraph.text.strip():
                    continue
                parts.append(paragraph.text.strip())
            return "\n".join(parts)
        except Exception as e:
            try:
                return cell.text.strip()
            except:
                return ""

    @classmethod
    def flatten_table(cls, table, depth=0, include_header=True):
        """
        Convert a docx Table into a structured list of lists with improved content preservation.
        """
        # Track merged cells to avoid duplicating content
        merged_cells = {}

        # First pass: detect all merged cells
        try:
            for i, row in enumerate(table.rows):
                for j, cell in enumerate(row.cells):
                    # Check for vMerge (vertical merge)
                    if hasattr(cell._element, "tcPr") and cell._element.tcPr:
                        vMerge = None
                        for elem in cell._element.tcPr:
                            if elem.tag.endswith("vMerge"):
                                vMerge = elem
                                vval = vMerge.get(qn("w:val")) if vMerge else None
                                # 'restart' indicates the start of a merged region
                                # If no value or 'continue', it's a continuation
                                if vval != "restart":
                                    # Find the starting cell
                                    for prev_i in range(i):
                                        if (prev_i, j) in merged_cells and merged_cells[(prev_i, j)]["active"]:
                                            merged_cells[(i, j)] = {"original": (prev_i, j), "active": False}
                                            break
                                else:
                                    merged_cells[(i, j)] = {"original": None, "active": True}
        except Exception as e:
            print(f"Warning: Error detecting merged cells: {e}")

        # First pass: collect all raw cell data
        raw_rows = []
        max_cols = max(len(row.cells) for row in table.rows)

        for i, row in enumerate(table.rows):
            if i == 0 and not include_header:
                continue

            row_cells = []
            for j, cell in enumerate(row.cells):
                # If this is a continuation of a merged cell, use empty string
                if (i, j) in merged_cells and merged_cells[(i, j)]["original"] is not None:
                    row_cells.append("")
                else:
                    cell_text = cls.extract_cell_text(cell)
                    row_cells.append(cell_text)

            # Ensure consistent column count
            row_cells += [""] * (max_cols - len(row_cells))
            raw_rows.append(row_cells)

        # Special handling for resume-style tables with hierarchical label structure
        if max_cols >= 2 and raw_rows:
            # Check if this looks like a resume table with main labels and sub-labels
            # Resume tables often have labels in first column that span multiple rows
            has_resume_structure = False

            # Check if most cells in first column are empty (suggesting merged cells)
            first_col = [row[0] for row in raw_rows]
            non_empty_count = sum(1 for cell in first_col if cell.strip())

            if non_empty_count > 0 and non_empty_count < len(first_col) * 0.6:
                has_resume_structure = True

            if has_resume_structure:
                # Process as a hierarchical structure
                hierarchical_rows = []
                current_main_label = None

                for row in raw_rows:
                    # If first cell has content, it's a main label
                    if row[0].strip():
                        current_main_label = row[0].strip()

                    # Process the rest of the row
                    if len(row) > 1 and row[1].strip():
                        # This is a sub-label with value
                        if len(row) > 2 and row[2].strip():
                            hierarchical_rows.append([current_main_label, row[1], row[2]])
                        else:
                            hierarchical_rows.append([current_main_label, row[1], ""])

                return hierarchical_rows

        # Default case: return processed rows
        return raw_rows

    @classmethod
    def detect_table_structure(cls, table):
        """
        Analyzes a table to determine its logical structure with special handling for resume tables.
        """
        # Count rows and columns
        row_count = len(table.rows)
        col_count = max(len(row.cells) for row in table.rows)

        if row_count == 0:
            return None

        # Extract all cell text
        cell_texts = []
        for row in table.rows:
            row_texts = []
            for cell in row.cells:
                row_texts.append(cls.extract_cell_text(cell))
            cell_texts.append(row_texts)

        # Handle resume-specific format (rightmost columns often have the actual data)
        if col_count >= 3:
            # Check if the rightmost column has valuable content (typical in resumes)
            rightmost_col_idx = col_count - 1
            rightmost_col = []
            for row in cell_texts:
                if rightmost_col_idx < len(row):
                    rightmost_col.append(row[rightmost_col_idx])

            # Check for email addresses and phone numbers in rightmost column
            has_contact_info = any(
                "@" in cell or any(c.isdigit() for c in cell) for cell in rightmost_col if cell  # Email pattern  # Phone numbers have digits
            )

            # If rightmost column has contact info, prioritize values from there
            if has_contact_info:
                # Process as a resume with values in rightmost column
                structured_data = []
                current_section = None

                for row in cell_texts:
                    if not row or not any(cell.strip() for cell in row):
                        continue

                    # First column might be section header
                    if row[0].strip():
                        current_section = row[0].strip()

                    # Look for label in middle columns
                    label = ""
                    for i in range(1, min(rightmost_col_idx, len(row))):
                        if row[i].strip() and row[i].strip() != row[0].strip():
                            label = row[i].strip()
                            break

                    # Get value from rightmost column
                    value = row[rightmost_col_idx].strip() if rightmost_col_idx < len(row) else ""

                    # Only add if we have a value or label
                    if value or label:
                        structured_data.append([current_section, label, value])

                if structured_data:
                    return structured_data

        # Continue with the existing duplicate column detection logic
        if col_count >= 2:
            # Identify and remove duplicate columns
            column_content = []
            for col_idx in range(col_count):
                col_values = []
                for row_idx in range(row_count):
                    if row_idx < len(cell_texts) and col_idx < len(cell_texts[row_idx]):
                        value = cell_texts[row_idx][col_idx].strip()
                        if value:
                            col_values.append(value)
                column_content.append(col_values)

            # Identify duplicate columns
            duplicate_cols = set()
            for i in range(col_count):
                if i in duplicate_cols:
                    continue

                for j in range(i + 1, col_count):
                    if j in duplicate_cols:
                        continue

                    # Skip if either column is empty
                    if not column_content[i] or not column_content[j]:
                        continue

                    # Compare content
                    matches = 0
                    total_comparisons = 0

                    # Compare values at same positions
                    for idx in range(min(len(column_content[i]), len(column_content[j]))):
                        if column_content[i][idx] and column_content[j][idx]:
                            total_comparisons += 1
                            if column_content[i][idx] == column_content[j][idx]:
                                matches += 1

                    # If highly similar, mark as duplicate
                    if total_comparisons > 0 and matches / total_comparisons > 0.8:
                        duplicate_cols.add(j)

            # If we found duplicate columns, create clean data
            if duplicate_cols:
                clean_columns = [i for i in range(col_count) if i not in duplicate_cols]
                clean_data = []
                for row in cell_texts:
                    clean_row = [row[col] if col < len(row) else "" for col in clean_columns]
                    clean_data.append(clean_row)

                cell_texts = clean_data
                col_count = len(clean_columns)

            # Now process the cleaned data
            # Check if first column contains section headers (same value repeats)
            first_col = [row[0] if row and len(row) > 0 else "" for row in cell_texts]
            first_col_values = [v for v in first_col if v.strip()]
            first_col_unique = set(first_col_values)

            # Resume-style tables often have fewer unique values in first column than rows
            if len(first_col_unique) < len(first_col_values) * 0.7:
                # Process as a resume table with sections
                structured_data = []
                current_section = None
                seen_label_values = set()  # Track to avoid duplicates

                for row in cell_texts:
                    if not row or not any(cell.strip() for cell in row):
                        continue

                    # Check if first column has content (section header)
                    if row[0].strip():
                        current_section = row[0].strip()

                    # For remaining columns, extract label-value pairs
                    if len(row) >= 3:  # Section + Label + Value
                        label = row[1].strip() if len(row) > 1 else ""
                        value = row[2].strip() if len(row) > 2 else ""

                        # Check for email addresses or phone numbers in all columns
                        # This helps catch important contact information that might be missed
                        for i in range(2, len(row)):
                            cell_content = row[i].strip()
                            if cell_content and ("@" in cell_content or any(c.isdigit() for c in cell_content)):
                                if label and not value:
                                    value = cell_content
                                elif label and value != cell_content:
                                    # Add as a separate entry
                                    pair_key = (label, cell_content)
                                    if pair_key not in seen_label_values:
                                        structured_data.append([current_section, label, cell_content])
                                        seen_label_values.add(pair_key)

                        # Add remaining values if present
                        if len(row) > 3:
                            extra_values = [v.strip() for v in row[3:] if v.strip()]
                            if extra_values:
                                if value:
                                    value = f"{value} {' '.join(extra_values)}"
                                else:
                                    value = " ".join(extra_values)

                        if (label or value) and (current_section, label, value) not in seen_label_values:
                            # Create row with section, label, value
                            structured_data.append([current_section, label, value])
                            seen_label_values.add((current_section, label, value))

                    elif len(row) == 2:  # Section + Content
                        content = row[1].strip() if len(row) > 1 else ""
                        if content and (current_section, "", content) not in seen_label_values:
                            structured_data.append([current_section, "", content])
                            seen_label_values.add((current_section, "", content))

                return structured_data

        # For other table types, return None to use the default handling
        return None

    @staticmethod
    def table_2d_to_string(rows_2d, align="left", spans=None):
        """
        Convert a 2D table list into well-structured text with improved handling of complex cases.
        Now handles merged cells and has better column width balancing.
        """
        if not rows_2d or not any(row for row in rows_2d):
            return "(Empty Table)"

        # Clean and normalize table data
        clean_rows = []
        max_cols = 0

        for row in rows_2d:
            if not any(cell for cell in row if cell and str(cell).strip()):
                continue  # Skip empty rows

            clean_row = [str(cell).strip() if cell is not None else "" for cell in row]
            max_cols = max(max_cols, len(clean_row))
            clean_rows.append(clean_row)

        if not clean_rows:
            return "(Empty Table)"

        # Pad rows to ensure consistent column count
        for i in range(len(clean_rows)):
            if len(clean_rows[i]) < max_cols:
                clean_rows[i] += [""] * (max_cols - len(clean_rows[i]))

        # Special case for resume-style hierarchical tables (3 columns with section, label, content)
        if max_cols == 3:
            # Check if first column contains section headers
            first_col = [row[0] for row in clean_rows if row[0].strip()]
            unique_sections = set(first_col)

            # If there are repeated sections (typical in resume format)
            if len(unique_sections) < len(first_col):
                result = []
                current_section = None
                used_pairs = set()  # Track label-value pairs to avoid duplicates

                for row in clean_rows:
                    section, label, value = row[0], row[1], row[2]

                    # Skip empty rows
                    if not section.strip() and not label.strip() and not value.strip():
                        continue

                    # Only print section when it changes
                    if section.strip() and section != current_section:
                        current_section = section
                        if result:  # Add spacing except for first section
                            result.append("")
                        result.append(f"{section}")

                    # Format label-value pair (on same line for better readability)
                    # Only include if we haven't seen this label-value pair before
                    pair = (label.strip(), value.strip())
                    if pair not in used_pairs:
                        if label.strip() and value.strip():
                            result.append(f"    {label.strip()}: {value.strip()}")
                            used_pairs.add(pair)
                        elif label.strip():
                            result.append(f"    {label.strip()}")
                        elif value.strip():
                            result.append(f"    {value.strip()}")

                # Remove leading blank line if present
                if result and not result[0].strip():
                    result.pop(0)

                return "\n".join(result)

        # Special case: Key-value tables (2 columns)
        if max_cols == 2:
            left_col = [row[0] for row in clean_rows]
            right_col = [row[1] for row in clean_rows]

            # Calculate average content length
            left_avg = sum(len(c) for c in left_col if c) / max(1, sum(1 for c in left_col if c))
            right_avg = sum(len(c) for c in right_col if c) / max(1, sum(1 for c in right_col if c))

            # If left column looks like labels (shorter than right column content)
            if left_avg < right_avg * 0.7:
                result = []
                current_section = None
                used_pairs = set()  # Track label-value pairs to avoid duplicates

                for row in clean_rows:
                    label, value = row[0], row[1]

                    # Skip empty rows
                    if not label.strip() and not value.strip():
                        continue

                    pair = (label.strip(), value.strip())
                    if pair not in used_pairs:
                        # If label might be a section header (all caps or short)
                        if label and (label.isupper() or len(label.split()) <= 2) and not current_section:
                            result.append(f"{label}")
                            current_section = label
                        elif label and value:
                            result.append(f"    {label}: {value}")
                            used_pairs.add(pair)
                        elif label:
                            result.append(f"    {label}")
                        elif value:
                            result.append(f"    {value}")

                return "\n".join(result)

        # Special case: Single column with paragraphs
        if max_cols == 1:
            if any(len(row[0]) > 80 for row in clean_rows):
                return "\n\n".join(row[0] for row in clean_rows if row[0])

        # Analyze column content for better width allocation and alignment
        col_types = []  # Will store 'numeric', 'short', or 'text' for each column
        col_widths = []
        MAX_TABLE_WIDTH = 120
        MIN_COL_WIDTH = 5

        for col_idx in range(max_cols):
            col_content = [row[col_idx] for row in clean_rows if col_idx < len(row)]

            # Detect if column is mostly numeric
            numeric_count = sum(1 for c in col_content if c and c.replace(".", "", 1).replace("-", "", 1).isdigit())
            is_numeric = numeric_count > len(col_content) * 0.7

            # Measure content length
            col_lengths = [len(c) for c in col_content if c]
            avg_len = sum(col_lengths) / max(1, len(col_lengths)) if col_lengths else 0
            max_len = max(col_lengths) if col_lengths else 0

            # Determine column type
            if is_numeric:
                col_types.append("numeric")
                # Numeric columns get just enough space plus small padding
                col_widths.append(min(max_len + 2, 15))
            elif avg_len < 15:
                col_types.append("short")
                # Short text gets a bit more than its average length
                col_widths.append(min(int(avg_len * 1.5) + 3, 20))
            else:
                col_types.append("text")
                # Regular text gets proportional space, but capped
                col_widths.append(min(int(avg_len * 1.2), 40))

        # Ensure minimum widths and adjust to total target width
        col_widths = [max(w, MIN_COL_WIDTH) for w in col_widths]
        total_width = sum(col_widths)

        # Scale widths to fit target if needed
        if total_width > MAX_TABLE_WIDTH:
            scale = MAX_TABLE_WIDTH / total_width
            col_widths = [max(int(w * scale), MIN_COL_WIDTH) for w in col_widths]

        # Helper function for smart text wrapping
        def wrap_text(text, width):
            if not text or len(text) <= width:
                return [text]

            # For very long words, force break
            words = []
            for word in text.split():
                if len(word) > width:
                    # Break long word into chunks
                    chunks = [word[i : i + width - 1] + "-" for i in range(0, len(word), width - 1)]
                    # Fix last chunk (no hyphen if it fits)
                    if chunks and len(chunks[-1]) <= width - 1:
                        chunks[-1] = chunks[-1][:-1]  # Remove hyphen
                    words.extend(chunks)
                else:
                    words.append(word)

            # Wrap text
            lines = []
            current = ""

            for word in words:
                if not current:
                    current = word
                elif len(current) + len(word) + 1 <= width:
                    current += " " + word
                else:
                    lines.append(current)
                    current = word

            if current:
                lines.append(current)

            return lines

        # Generate markdown table with proper alignment
        table_lines = []

        # Process header row
        header_wrapped = []
        for i, cell in enumerate(clean_rows[0]):
            wrapped = wrap_text(cell, col_widths[i])
            header_wrapped.append(wrapped)

        # Get max header height
        header_height = max(len(cell) for cell in header_wrapped)

        # Generate header lines
        for line_idx in range(header_height):
            row_cells = []
            for i, wrapped_cell in enumerate(header_wrapped):
                cell_text = wrapped_cell[line_idx] if line_idx < len(wrapped_cell) else ""
                # Headers are always left-aligned
                row_cells.append(cell_text.ljust(col_widths[i]))
            table_lines.append("| " + " | ".join(row_cells) + " |")

        # Add separator with alignment indicators
        separators = []
        for i, col_type in enumerate(col_types):
            width = col_widths[i]
            if col_type == "numeric":
                # Right align numbers
                separators.append("-" * (width - 1) + ":")
            else:
                # Left align text
                separators.append(":" + "-" * (width - 1))
        table_lines.append("| " + " | ".join(separators) + " |")

        # Process data rows
        for row_idx in range(1, len(clean_rows)):
            row = clean_rows[row_idx]

            # Wrap each cell
            row_wrapped = []
            for i, cell in enumerate(row):
                if i >= len(col_widths):
                    break
                wrapped = wrap_text(cell, col_widths[i])
                row_wrapped.append(wrapped)

            # Fill in any missing cells
            while len(row_wrapped) < max_cols:
                row_wrapped.append([""])

            # Get max row height
            row_height = max(len(cell) for cell in row_wrapped)

            # Generate row lines
            for line_idx in range(row_height):
                row_cells = []
                for i, wrapped_cell in enumerate(row_wrapped):
                    cell_text = wrapped_cell[line_idx] if line_idx < len(wrapped_cell) else ""

                    # Align based on content type
                    if col_types[i] == "numeric":
                        row_cells.append(cell_text.rjust(col_widths[i]))  # Right-align numbers
                    else:
                        row_cells.append(cell_text.ljust(col_widths[i]))  # Left-align text

                table_lines.append("| " + " | ".join(row_cells) + " |")

        return "\n".join(table_lines)

    @staticmethod
    def content_similarity(text1, text2):
        """text similarity comparison using thefuzz"""
        if not text1 or not text2:
            return 0

        # Normalize texts
        t1 = text1.lower().strip()
        t2 = text2.lower().strip()

        if t1 == t2:
            return 1.0

        # Length comparison (quick filter)
        len_ratio = min(len(t1), len(t2)) / max(len(t1), len(t2))
        if len_ratio < 0.5:
            return len_ratio * 0.3

        # Use multiple fuzzy matching algorithms for better results
        ratio = fuzz.ratio(t1, t2) / 100.0
        token_sort = fuzz.token_sort_ratio(t1, t2) / 100.0  # Handles word reordering

        # For longer texts, also use token_set_ratio which ignores duplicated words
        if len(t1) > 50 and len(t2) > 50:
            token_set = fuzz.token_set_ratio(t1, t2) / 100.0
            # Weighted average of different metrics
            return (0.2 * len_ratio) + (0.3 * ratio) + (0.25 * token_sort) + (0.25 * token_set)
        else:
            # For shorter texts
            return (0.3 * len_ratio) + (0.35 * ratio) + (0.35 * token_sort)

    @classmethod
    def markdown_table_from_docx_table(cls, docx_table):
        """
        Convert a docx Table into a markdown-formatted string with robust error handling
        and better handling of complex tables.
        """
        try:
            # Check for empty or invalid tables
            if not docx_table.rows or len(docx_table.rows) == 0:
                return "(Empty Table)"

            # Track merged cells
            spans = {}
            try:
                for i, row in enumerate(docx_table.rows):
                    for j, cell in enumerate(row.cells):
                        # Check for vertical and horizontal spans
                        if hasattr(cell._element, "tcPr") and cell._element.tcPr:
                            # Get vMerge and gridSpan directly without using xpath
                            vMerge = None
                            gridSpan = None

                            for elem in cell._element.tcPr:
                                if elem.tag.endswith("vMerge"):
                                    vMerge = elem
                                elif elem.tag.endswith("gridSpan"):
                                    gridSpan = elem

                            if vMerge or gridSpan:
                                spans[(i, j)] = {
                                    "vMerge": vMerge.get(qn("w:val")) if vMerge and vMerge.get(qn("w:val")) else "continue" if vMerge else None,
                                    "gridSpan": int(gridSpan.get(qn("w:val"))) if gridSpan else 1,
                                }
            except Exception as e:
                # Continue even if span detection fails
                print(f"Warning: Could not detect cell spans in table: {e}")

            # Store original content for comparison
            original_content = []
            for row in docx_table.rows:
                row_content = []
                for cell in row.cells:
                    row_content.append(cls.extract_cell_text(cell))
                original_content.append(row_content)

            # Try to detect table structure
            detected_structure = cls.detect_table_structure(docx_table)
            if detected_structure:
                return cls.table_2d_to_string(detected_structure, spans=spans)

            # Try advanced table flattening
            flattened = cls.flatten_table(docx_table, include_header=True)

            # If flattening didn't lose too much content, use it
            original_char_count = sum(len(str(cell)) for row in original_content for cell in row if cell)
            flattened_char_count = sum(len(str(cell)) for row in flattened for cell in row if cell)

            # If we preserved at least 85% of content, use flattened version
            if flattened_char_count >= 0.85 * original_char_count:
                return cls.table_2d_to_string(flattened, spans=spans)

            # Fall back to minimal processing of original content
            # Just clean up any obvious duplicate cells
            cleaned_original = []
            for row in original_content:
                if not row:
                    continue

                # Remove duplicate adjacent cells
                cleaned_row = []
                prev_cell = None
                for cell in row:
                    if cell != prev_cell or not cell:  # Keep empty cells and unique content
                        cleaned_row.append(cell)
                        prev_cell = cell

                cleaned_original.append(cleaned_row)

            return cls.table_2d_to_string(cleaned_original, spans=spans)

        except Exception as e:
            # Fallback to very simple table extraction
            try:
                rows = []
                for row in docx_table.rows:
                    row_data = []
                    for cell in row.cells:
                        try:
                            row_data.append(cell.text.strip())
                        except:
                            row_data.append("")
                    rows.append(row_data)
                return cls.table_2d_to_string(rows)
            except Exception as inner_e:
                return f"(Could not process table: {str(inner_e)})"

    @staticmethod
    def extract_footnotes(doc):
        """Extract footnotes and footnote references with enhanced structure."""
        footnotes = {}
        footnote_refs = {}

        # Extract footnotes from the document
        if hasattr(doc.part, "footnotes") and doc.part.footnotes:
            for idx, footnote in enumerate(
                doc.part.footnotes.element.findall(".//w:footnote", namespaces={"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"})
            ):

                # Skip separator and continuation separator footnotes (ids -1 and 0)
                footnote_id = footnote.get(qn("w:id"))
                if footnote_id and int(footnote_id) < 1:
                    continue

                # Collect all text from the footnote
                note_parts = [
                    node.text
                    for node in footnote.findall(".//w:t", namespaces={"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"})
                    if node.text
                ]
                note_text = " ".join(note_parts).strip()

                if note_text:  # Only store non-empty footnotes
                    footnotes[idx + 1] = note_text

        # Capture actual footnote references in the document
        for para in doc.paragraphs:
            for run in para.runs:
                for ref in run._element.findall(
                    ".//w:footnoteReference", namespaces={"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
                ):

                    ref_id = ref.get(qn("w:id"))
                    if ref_id and int(ref_id) > 0:  # Skip separator footnotes
                        footnote_refs[run.text.strip()] = f"[^{ref_id}]"

        return footnotes, footnote_refs

    @staticmethod
    def handle_lists(para):
        """Detects and formats bulleted and numbered lists dynamically."""
        # Extended bullet character list
        bullet_chars = {"-", "*", "•", "▪", "▶", "◾", "◼", "✔", "☑", "✓", "→", "⇒"}

        text = para.text.strip()
        if not text:
            return None

        # Check for DOCX auto-numbered lists
        has_numbering = para._p.find(qn("w:numPr")) is not None

        # Check for manual bullet lists
        first_char = text[0] if text else ""
        is_bullet_list = first_char in bullet_chars

        # Check for manual numbered lists (e.g., "1.", "2.", etc.)
        is_numbered_list = bool(re.match(r"^\d+[\.\)]", text))

        if has_numbering or is_bullet_list or is_numbered_list:
            # Determine indentation level
            left_indent = para.paragraph_format.left_indent
            indent_level = 0 if left_indent is None else min(3, int(left_indent.pt / 18))  # Cap at 3 levels
            indent_spaces = "  " * indent_level  # Two spaces per indent level

            # Format based on list type
            if is_bullet_list:
                # Remove the bullet character and any leading whitespace
                clean_text = text[1:].lstrip()
                return f"{indent_spaces}- {clean_text}"
            elif is_numbered_list:
                # Extract the number and text
                match = re.match(r"^(\d+)[\.\)](.*)$", text)
                if match:
                    number, content = match.groups()
                    clean_text = content.lstrip()
                    return f"{indent_spaces}{number}. {clean_text}"
            else:  # Auto-numbered list from Word
                # For auto-numbered lists, use a dash as we can't reliably get the number
                clean_text = text.strip()
                return f"{indent_spaces}- {clean_text}"

        return None  # Not a list item

    @classmethod
    def extract_docx_structure(cls, docx_file_path: str) -> str:
        """
        Extracts structured text from a DOCX file with robust error handling
        and improved handling of tables, captions, and duplicate detection.

        Returns: str on successful parsing
                 False on media related errors
                 None for other broad spectrum errors
        """
        # Load DOCX file
        try:
            doc = Document(docx_file_path)
        except Exception as e:
            error_message = str(e).lower()

            # Comprehensive list of media-related error indicators
            media_error_indicators = [
                # General media terms
                "image",
                "picture",
                "graphic",
                "media",
                "visual",
                "illustration",
                "photo",
                "photograph",
                "artwork",
                "icon",
                "logo",
                "diagram",
                "figure",
                "chart",
                "embed",
                "embedded",
                "object",
                # Image formats
                "png",
                "jpeg",
                "jpg",
                "tiff",
                "bmp",
                "gif",
                "svg",
                "webp",
                "heic",
                "heif",
                "raw",
                "psd",
                "ai",
                "eps",
                "pdf",
                "ico",
                "exif",
                "dib",
                "pcx",
                "tga",
                "jp2",
                "j2k",
                "jpx",
                "jpm",
                "mj2",
                "jxr",
                "hdp",
                "wdp",
                "avif",
                # Video formats
                "video",
                "movie",
                "clip",
                "animation",
                "motion",
                "mp4",
                "avi",
                "mov",
                "wmv",
                "flv",
                "mkv",
                "webm",
                "m4v",
                "3gp",
                "mpeg",
                "mpg",
                "asf",
                "rm",
                "rmvb",
                "vob",
                "ogv",
                "ogg",
                "dv",
                "qt",
                "yuv",
                "m2v",
                "m4p",
                # Animated formats
                "gif",
                "animated",
                "animation",
                "apng",
                "mng",
                "webp",
                "flc",
                "fli",
                # OLE objects and embedding terms
                "ole",
                "object linking",
                "embedded",
                "object",
                "activex",
                "ole2",
                # Document relationship terms
                "relationship",
                "part",
                "content type",
                "unsupported content",
                "corrupt",
                "malformed",
                "invalid",
                "broken",
                "damaged",
                # Microsoft Office specific
                "shape",
                "smartart",
                "wordart",
                "drawing",
                "canvas",
                "vml",
                "office art",
                "drawingml",
                "pictureid",
                "imageid",
                "mediaid",
            ]

            # Check if any media-related terms appear in the error message
            if any(indicator in error_message for indicator in media_error_indicators):
                print(f"Media-related error detected: {str(e)}")
                return False

            # For all other errors
            print(f"Non-media error opening document: {str(e)}")
            return None

        output = []

        # Enhanced tracking systems
        table_content = {}  # Store table content with position info
        processed_content = set()  # Track processed fingerprints

        # Helper function to detect table captions
        def is_table_caption(paragraph, next_block_is_table):
            if not paragraph.text.strip():
                return False

            text = paragraph.text.strip().lower()

            # Check for explicit table references
            has_table_ref = text.startswith("table") or "table " in text or re.search(r"tab\.?\s+\d", text)

            # Check for caption-like formatting
            has_caption_format = ":" in text or bool(re.match(r"^(figure|fig|table|tab)\.?\s*\d", text))

            # If it looks like a caption and the next block is a table, it's likely a table caption
            return next_block_is_table and (has_table_ref or has_caption_format)

        # First pass - collect all tables and their content
        try:
            blocks = list(cls.iter_block_items(doc))

            # Process tables to collect their content
            for i, block in enumerate(blocks):
                if isinstance(block, Table):
                    table_id = id(block._element)

                    # Process all cells
                    for row_idx, row in enumerate(block.rows):
                        for col_idx, cell in enumerate(row.cells):
                            cell_text = cls.extract_cell_text(cell)
                            if not cell_text or len(cell_text) < 15:
                                continue

                            # Create fingerprint
                            fingerprint = cls.get_content_fingerprint(cell_text)
                            if fingerprint:
                                table_content[fingerprint] = {"text": cell_text, "table_id": table_id, "position": (row_idx, col_idx)}

                            # Also process paragraphs within cell
                            for para in cell_text.split("\n"):
                                if para and len(para) > 20:
                                    para_fp = cls.get_content_fingerprint(para)
                                    if para_fp:
                                        table_content[para_fp] = {
                                            "text": para,
                                            "table_id": table_id,
                                            "position": (row_idx, col_idx),
                                            "is_paragraph": True,
                                        }
        except Exception as e:
            print(f"Warning in first pass: {e}")

        # Extract footnotes
        try:
            footnotes, footnote_refs = cls.extract_footnotes(doc)
        except Exception as e:
            print(f"Warning extracting footnotes: {e}")
            footnotes, footnote_refs = {}, {}

        # Track state for better formatting
        in_list = False
        list_items = []
        prev_was_header = False
        prev_was_table = False
        prev_was_caption = False

        # Second pass - process document content with caption awareness
        try:
            for i, block in enumerate(blocks):
                # Check if the next block is a table (for caption detection)
                next_is_table = i + 1 < len(blocks) and isinstance(blocks[i + 1], Table)

                if isinstance(block, Paragraph):
                    # Check if this is a table caption
                    if is_table_caption(block, next_is_table):
                        # Format caption specially
                        caption_text = block.text.strip()
                        output.append(f"**{caption_text}**")
                        prev_was_caption = True
                        continue

                    # Skip paragraphs in tables
                    if hasattr(block, "_in_table_context") and block._in_table_context:
                        continue

                    para_text = block.text.strip()

                    # Skip empty paragraphs unless needed for spacing
                    if not para_text:
                        if output and output[-1] and not prev_was_table and not prev_was_caption:
                            output.append("")
                        continue

                    # Check for duplicate content using thefuzz
                    if len(para_text) > 15:
                        para_fp = cls.get_content_fingerprint(para_text)

                        # Skip if we've already processed this exact content
                        if para_fp in processed_content:
                            continue

                        # Check if similar to table content
                        is_duplicate = False
                        for table_fp, table_info in table_content.items():
                            similarity = cls.content_similarity(para_text, table_info["text"])
                            # Skip if very similar
                            if similarity > 0.85:
                                is_duplicate = True
                                break

                        if is_duplicate:
                            processed_content.add(para_fp)
                            continue

                    # Check if it's a list item
                    list_item = cls.handle_lists(block)
                    if list_item:
                        if not in_list:
                            if output and output[-1] and not prev_was_table and not prev_was_header and not prev_was_caption:
                                output.append("")
                            in_list = True
                        list_items.append(list_item)
                        continue
                    elif in_list:
                        output.extend(list_items)
                        output.append("")
                        list_items = []
                        in_list = False

                    # Check if it's a header
                    if cls.is_header(block):
                        if output and output[-1] and not prev_was_header and not prev_was_table and not prev_was_caption:
                            output.append("")

                        if para_text.isupper() and len(para_text.split()) <= 3:
                            output.append(f"### {para_text} ###")
                        else:
                            output.append(f"### {para_text}")

                        if len(para_text) > 10:
                            processed_content.add(cls.get_content_fingerprint(para_text))

                        prev_was_header = True
                        prev_was_table = False
                        prev_was_caption = False
                    else:
                        # Regular paragraph
                        if (prev_was_table or prev_was_caption) and output and output[-1]:
                            if output[-1].strip():
                                output.append("")

                        output.append(para_text)

                        if len(para_text) > 15:
                            processed_content.add(cls.get_content_fingerprint(para_text))

                        prev_was_header = False
                        prev_was_table = False
                        prev_was_caption = False

                elif isinstance(block, Table):
                    if output and output[-1] and not output[-1].isspace() and not prev_was_table and not prev_was_caption:
                        output.append("")

                    # Process table with robust error handling
                    table_md = cls.markdown_table_from_docx_table(block)
                    output.append(table_md)

                    # Add spacing after table
                    output.append("")

                    prev_was_table = True
                    prev_was_header = False
                    prev_was_caption = False

        except Exception as e:
            output.append(f"\nError processing document: {str(e)}")

        # Handle any remaining list items
        if list_items:
            output.extend(list_items)

        # Add footnotes if they exist
        if footnotes:
            output.append("\nFootnotes:")
            for idx, footnote in sorted(footnotes.items()):
                output.append(f"[{idx}] {footnote}")

        # Clean up output - remove consecutive blank lines
        cleaned_output = []
        for i, line in enumerate(output):
            if i > 0 and not line and not output[i - 1]:
                continue  # Skip consecutive blank lines
            cleaned_output.append(line)

        return "\n".join(cleaned_output)

    @staticmethod
    def is_header(para):
        """
        Dynamically determines if a paragraph is a header based on formatting characteristics.
        """
        # Skip empty paragraphs
        if not para.text.strip() or not para.runs:
            return False

        # Get paragraph style name
        style_name = para.style.name.lower() if para.style and hasattr(para.style, "name") else ""

        # Check if explicitly styled as a heading
        if "heading" in style_name or style_name.startswith("h"):
            return True

        # Exclude paragraphs explicitly styled as lists
        if any(x in style_name for x in ["list", "bullet", "number"]):
            return False

        # Check if this paragraph is part of a list
        if para._p.find(qn("w:numPr")) is not None:
            return False

        # Check if this paragraph is part of a table (additional check)
        if para._p.find(qn("w:tbl")) is not None or para._p.getparent().tag == qn("w:tc"):
            return False

        # Get paragraph text
        text = para.text.strip()

        # Very short text that's just a single word is likely a table cell, not a header
        # (unless it has strong header formatting)
        if len(text.split()) == 1 and not text.isupper():
            # Check if it has strong header formatting before deciding
            if not para.style.name.lower().startswith("heading"):
                return False

        # Check for formatting characteristics that suggest a header
        is_bold = all(run.bold for run in para.runs if run.text.strip())
        is_larger_font = False

        # Check for larger font size compared to default
        try:
            default_size = 11  # Typical default size
            if para.runs and hasattr(para.runs[0], "font") and para.runs[0].font.size:
                font_size = para.runs[0].font.size.pt
                is_larger_font = font_size > default_size + 2  # At least 2pt larger
        except:
            pass

        # Check for ALL CAPS text (strong indicator of headers in many documents)
        is_all_caps = text.isupper() and len(text) > 3

        # Check if text ends with typical non-header punctuation
        ends_with_sentence_punct = text.endswith((".", "?", "!", ":", ";"))

        # Short text with header-like formatting
        is_short = len(text.split()) <= 8

        # Apply heuristics to determine if this is a header
        if (is_all_caps and is_short) or (is_bold and is_short and not ends_with_sentence_punct):
            return True

        if is_larger_font and is_short and not ends_with_sentence_punct:
            return True

        # Check for numbered headers like "1.2 Section Title"
        if re.match(r"^\d+(\.\d+)*\s+\S", text) and is_short and (is_bold or is_larger_font):
            return True

        return False

    @classmethod
    # wrapper to make extract_docx_structure async without a rewrite
    async def extract_docx_structure_async(cls, docx_file_path: str) -> str:
        import asyncio

        return await asyncio.to_thread(cls.extract_docx_structure, docx_file_path)


if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Extract text from a DOCX file.")
    parser.add_argument("docx_file", help="The path to the DOCX file")

    args = parser.parse_args()

    # Call the function with the arguments
    print(DocxExtractor.extract_docx_structure(args.docx_file))
